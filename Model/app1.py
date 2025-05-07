# Standard library imports
import os
import json
import threading
import logging
from flask import Flask, request, jsonify, send_from_directory
from flask_cors import CORS
from openai import OpenAI
from langchain_community.document_loaders import PyPDFLoader, Docx2txtLoader, TextLoader, CSVLoader, UnstructuredExcelLoader
from langchain_community.vectorstores import FAISS
from langchain_core.embeddings import Embeddings
from langchain_core.prompts import ChatPromptTemplate
from langchain_core.output_parsers import StrOutputParser
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain.document_transformers import LongContextReorder
from langchain.schema.runnable import RunnableLambda
from langchain.schema.runnable.passthrough import RunnableAssign
from langchain_nvidia_ai_endpoints import ChatNVIDIA, NVIDIAEmbeddings
import numpy as np
from operator import itemgetter
from functools import partial
from pprint import pprint

# Configure logging
logging.basicConfig(level=logging.DEBUG)

# Clear SSL issues
if 'SSL_CERT_FILE' in os.environ:
    os.environ.pop('SSL_CERT_FILE')

# Replace with direct API key assignment
NVIDIA_API_KEY = "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"
os.environ["NVIDIA_API_KEY"] = NVIDIA_API_KEY
print(f"NVIDIA API Key configured: {NVIDIA_API_KEY[:10]}...")

# Configuration
BASE_URL = "https://integrate.api.nvidia.com/v1"
MODEL = "mistralai/mixtral-8x7b-instruct-v0.1"
MAX_TOKENS = 1024
TEMPERATURE = 0.5

# Initialize Flask app
app = Flask(__name__, static_folder='UI')
CORS(app, resources={r"/*": {"origins": "*"}})

# Create the OpenAI client with the correct configuration
client = OpenAI(
    base_url=BASE_URL,
    api_key=NVIDIA_API_KEY
)

# Track embedding status
embeddings_ready = False
vectors = None
embeddings = None
docstore = None
convstore = None

# Initialize NVIDIA models
embedder = NVIDIAEmbeddings(model="nvidia/nv-embed-v1")
instruct_llm = ChatNVIDIA(model="mistralai/mixtral-8x7b-instruct-v0.1")

# Utility functions for RAG
def RPrint(preface=""):
    """Simple passthrough 'prints, then returns' chain"""
    def print_and_return(x, preface=preface):
        if preface:
            print(preface, end="")
        print(x)
        return x
    return RunnableLambda(print_and_return)

def docs2str(docs, title="Document"):
    """Useful utility for making chunks into context string."""
    out_str = ""
    for doc in docs:
        doc_name = getattr(doc, 'metadata', {}).get('Title', title)
        if doc_name:
            out_str += f"[Quote from {doc_name}] "
        out_str += getattr(doc, 'page_content', str(doc)) + "\n"
    return out_str

# Reorder to place longer documents in the center
long_reorder = RunnableLambda(LongContextReorder().transform_documents)

# Function to save conversation history
def save_memory_and_get_output(d, vstore):
    """Accepts 'input'/'output' dictionary and saves to conversation store"""
    vstore.add_texts([f"User said {d.get('input')}", f"Agent said {d.get('output')}"])
    return d.get('output')

# Function to initialize document embeddings
def initialize_embeddings():
    global embeddings_ready, vectors, embeddings, docstore, convstore

    print("Starting document embeddings initialization...")

    try:
        # Make sure Data directory exists
        data_path = os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))), 'Data')
        print(f"Looking for Data folder at: {data_path}")

        # Create an empty list to store all documents
        documents = []

        # Process each file individually by type
        for root, _, files in os.walk(data_path):
            for file in files:
                file_path = os.path.join(root, file)
                try:
                    # Skip JavaScript files and other non-document types
                    if file.endswith('.js'):
                        continue

                    # Use appropriate loader based on file extension
                    if file.lower().endswith('.pdf'):
                        from langchain_community.document_loaders import PyPDFLoader
                        loader = PyPDFLoader(file_path)
                        documents.extend(loader.load())
                        print(f"Loaded PDF: {file}")
                    elif file.lower().endswith('.docx'):
                        # Alternative method using python-docx instead of docx2txt
                        try:
                            import docx
                            doc = docx.Document(file_path)
                            text_content = "\n\n".join([para.text for para in doc.paragraphs if para.text.strip()])

                            # Also extract tables if present
                            for table in doc.tables:
                                for row in table.rows:
                                    row_text = " | ".join([cell.text for cell in row.cells])
                                    text_content += "\n" + row_text

                            from langchain_core.documents import Document
                            documents.append(Document(
                                page_content=text_content,
                                metadata={"source": file_path, "title": file}
                            ))
                            print(f"Loaded DOCX: {file}")
                        except ImportError:
                            print(f"Error loading {file}: Missing python-docx module")
                            print("Install with: pip install python-docx")
                            continue
                    elif file.lower().endswith('.txt'):
                        from langchain_community.document_loaders import TextLoader
                        loader = TextLoader(file_path)
                        documents.extend(loader.load())
                        print(f"Loaded TXT: {file}")
                    elif file.lower().endswith('.xlsx') or file.lower().endswith('.xls'):
                        # Convert Excel data to text documents
                        try:
                            import pandas as pd
                            excel_file = pd.ExcelFile(file_path)
                            for sheet_name in excel_file.sheet_names:
                                df = pd.read_excel(file_path, sheet_name=sheet_name)
                                # Convert DataFrame to string representation
                                text_content = f"Excel File: {file}, Sheet: {sheet_name}\n\n"

                                # Add column headers
                                text_content += "| " + " | ".join(str(col) for col in df.columns) + " |\n"
                                text_content += "| " + " | ".join("-" * len(str(col)) for col in df.columns) + " |\n"

                                # Add rows
                                for _, row in df.iterrows():
                                    text_content += "| " + " | ".join(str(cell) for cell in row) + " |\n"

                                # Create document
                                from langchain_core.documents import Document
                                doc = Document(
                                    page_content=text_content,
                                    metadata={"source": file_path, "sheet": sheet_name}
                                )
                                documents.append(doc)
                            print(f"Loaded Excel: {file}")
                        except Exception as excel_error:
                            print(f"Error processing Excel file {file}: {str(excel_error)}")
                    else:
                        print(f"Skipped unsupported file: {file}")
                except Exception as e:
                    print(f"Error loading {file}: {str(e)}")
                    continue

        if documents:
            print(f"Number of Documents Retrieved: {len(documents)}")

            # Split documents
            text_splitter = RecursiveCharacterTextSplitter(
                chunk_size=1200,
                chunk_overlap=100,
                separators=["\n\n", "\n", ".", ";", ",", " ", ""]
            )

            docs_split = text_splitter.split_documents(documents)
            print(f"Total chunks after splitting: {len(docs_split)}")

            # Create FAISS index from document chunks
            global docstore
            docstore = FAISS.from_documents(docs_split, embedder)
            print(f"Created document store with {len(docs_split)} chunks")

            # Initialize conversation history
            conversation = [
                "User asked about document analysis",
                "Agent explained how documents are processed",
                "User asked about embeddings",
                "Agent explained vector embeddings for text"
            ]

            # Create conversation vector store
            global convstore
            convstore = FAISS.from_texts(conversation, embedding=embedder)
            print(f"Created conversation store with {len(conversation)} entries")

            embeddings_ready = True
            print("Document embeddings initialization completed successfully")

        else:
            print("No documents were loaded from the directory")

    except Exception as e:
        print(f"Error initializing embeddings: {str(e)}")
        return

# Start embeddings initialization in a background thread
threading.Thread(target=initialize_embeddings).start()

# Routes
@app.route('/')
def index():
    return send_from_directory('../UI', 'chatbot.html')

@app.route('/api/status', methods=['GET'])
def status():
    return jsonify({
        'nim_api': 'configured',
        'embeddings_ready': embeddings_ready
    })

@app.route('/api/query-documents', methods=['POST'])
def query_documents():
    global docstore, embeddings, embeddings_ready

    data = request.json
    query = data.get('query')

    if not query:
        return jsonify({
            'status': 'error',
            'message': 'Query is required'
        }), 400

    if not embeddings_ready:
        return jsonify({
            'status': 'error',
            'message': 'Document embeddings are not ready yet. Please try again later.'
        }), 503

    try:
        # Perform search
        results = docstore.similarity_search_with_score(query, k=5)

        # Format results
        formatted_results = []
        for doc, score in results:
            formatted_results.append({
                "content": doc.page_content,
                "metadata": doc.metadata,
                "score": float(score)
            })

        return jsonify({
            'status': 'success',
            'data': {
                'results': formatted_results
            }
        })

    except Exception as e:
        return jsonify({
            'status': 'error',
            'message': str(e)
        }), 500

@app.route('/api/chat', methods=['POST'])
def chat():
    """Process chat messages with NVIDIA API"""
    try:
        print("Received chat request")
        data = request.json

        if not data or not isinstance(data.get('messages'), list):
            return jsonify({
                "status": "error",
                "message": "Invalid request: messages array is required"
            }), 400

        messages = data.get('messages', [])

        # Make API request
        response = client.chat.completions.create(
            model=MODEL,
            messages=messages,
            temperature=TEMPERATURE,
            max_tokens=MAX_TOKENS,
            stream=False
        )

        response_text = response.choices[0].message.content

        return jsonify({
            "status": "success",
            "message": response_text
        })
    except Exception as e:
        print(f"Error in chat endpoint: {str(e)}")
        return jsonify({
            "status": "error",
            "message": f"Server error: {str(e)}"
        }), 500

@app.route('/api/basic_chat', methods=['POST'])
def basic_chat():
    """Basic chat endpoint without RAG"""
    try:
        data = request.json
        query = data.get('query', '')

        if not query:
            return jsonify({
                "status": "error",
                "message": "Query parameter is required"
            }), 400

        # Use the NVIDIA ChatNVIDIA model directly
        response = instruct_llm.invoke(query)
        response_text = response.content

        return jsonify({
            "status": "success",
            "message": response_text
        })
    except Exception as e:
        print(f"Error in basic_chat endpoint: {str(e)}")
        return jsonify({
            "status": "error",
            "message": f"Server error: {str(e)}"
        }), 500

@app.route('/api/rag_chat', methods=['POST'])
def rag_chat():
    """RAG-enhanced chat endpoint"""
    global docstore, convstore, embeddings_ready

    print("RAG chat endpoint called")

    try:
        data = request.json
        query = data.get('query', '')

        print(f"RAG query: {query}")
        print(f"Embeddings ready: {embeddings_ready}")
        print(f"Docstore exists: {docstore is not None}")

        if not query:
            return jsonify({
                "status": "error",
                "message": "Query parameter is required"
            }), 400

        # Check if embeddings are ready
        if not embeddings_ready or docstore is None:
            print("WARNING: Document embeddings not ready yet, falling back to basic chat")

            # Fall back to basic chat if embeddings aren't ready
            response = instruct_llm.invoke(query)
            return jsonify({
                "status": "success",
                "message": response.content,
                "debug_info": {
                    "used_rag": False,
                    "reason": "Embeddings not ready",
                    "embeddings_ready": embeddings_ready,
                    "docstore_exists": docstore is not None
                }
            })

        try:
            # SIMPLIFIED RAG APPROACH - bypass complex LangChain chains
            # 1. Get relevant documents
            print("Retrieving relevant documents")
            results = docstore.similarity_search_with_score(query, k=5)

            # 2. Format context from documents
            context = ""
            for doc, score in results:
                source = doc.metadata.get('source', 'Unknown').split('/')[-1]
                context += f"\n[From {source}]: {doc.page_content}\n"

            print(f"Retrieved context (first 100 chars): {context[:100]}...")

            # 3. Create prompt with context
            prompt = f"""Answer the following question using the information from the documents below.
If the documents don't contain the necessary information, say so and provide your best answer.

DOCUMENTS:
{context}

QUESTION: {query}

ANSWER:"""

            # 4. Send to LLM
            print("Sending to LLM")
            response = instruct_llm.invoke(prompt)

            return jsonify({
                "status": "success",
                "message": response.content,
                "debug_info": {
                    "used_rag": True,
                    "context_length": len(context),
                    "retrieved_docs": len(results)
                }
            })

        except Exception as e:
            print(f"Error in RAG processing: {str(e)}")
            import traceback
            traceback.print_exc()

            # Fall back to regular chat if RAG fails
            response = instruct_llm.invoke(query)
            return jsonify({
                "status": "fallback",
                "message": response.content,
                "debug_info": {
                    "rag_error": str(e),
                    "used_fallback": True
                }
            })

    except Exception as e:
        print(f"Error in rag_chat endpoint: {str(e)}")
        import traceback
        traceback.print_exc()
        return jsonify({
            "status": "error",
            "message": f"Server error: {str(e)}",
            "error": str(e)
        }), 500

@app.route('/api/generator', methods=['POST'])
def generator():
    """Content generation endpoint"""
    try:
        data = request.json
        prompt = data.get('prompt', '')

        if not prompt:
            return jsonify({
                "status": "error",
                "message": "Prompt parameter is required"
            }), 400

        # Create generation prompt
        generation_prompt = ChatPromptTemplate.from_template(
            "Generate creative content based on the following prompt: {prompt}\n\n"
            "Be detailed, engaging, and original in your response."
        )

        # Create generation chain
        generation_chain = (
            {'prompt': (lambda x: x)}
            | generation_prompt
            | instruct_llm
            | StrOutputParser()
        )

        # Get response
        response_text = generation_chain.invoke(prompt)

        return jsonify({
            "status": "success",
            "message": response_text
        })
    except Exception as e:
        print(f"Error in generator endpoint: {str(e)}")
        return jsonify({
            "status": "error",
            "message": f"Server error: {str(e)}"
        }), 500

@app.route('/api/retriever', methods=['POST'])
def retriever():
    """Document retrieval endpoint"""
    global docstore, embeddings_ready

    try:
        data = request.json
        query = data.get('query', '')

        if not query:
            return jsonify({
                "status": "error",
                "message": "Query parameter is required"
            }), 400

        if not embeddings_ready:
            return jsonify({
                "status": "error",
                "message": "Document embeddings are not ready yet. Please try again later."
            }), 503

        # Perform search
        results = docstore.similarity_search_with_score(query, k=5)

        # Format results
        formatted_results = []
        for doc, score in results:
            formatted_results.append({
                "content": doc.page_content,
                "metadata": doc.metadata,
                "score": float(score)
            })

        return jsonify({
            "status": "success",
            "data": formatted_results
        })
    except Exception as e:
        print(f"Error in retriever endpoint: {str(e)}")
        return jsonify({
            "status": "error",
            "message": f"Server error: {str(e)}"
        }), 500

@app.route('/<path:filename>')
def serve_static(filename):
    # Try UI folder first
    ui_path = os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))), 'UI')
    if os.path.exists(os.path.join(ui_path, filename)):
        return send_from_directory('../UI', filename)

    # Then try Data folder
    data_path = os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))), 'Data')
    if os.path.exists(os.path.join(data_path, filename)):
        return send_from_directory('../Data', filename)

    return "File not found", 404

@app.route('/api/nim/status', methods=['GET'])
def nim_status():
    """Endpoint to check NVIDIA NIM API status"""
    try:
        # Test the connection with a minimal query
        response = client.chat.completions.create(
            model=MODEL,
            messages=[{"role": "user", "content": "Hi"}],
            max_tokens=5,
            stream=False
        )

        # Connection successful
        return jsonify({
            "connected": True,
            "status": "ok",
            "model": MODEL,
            "base_url": BASE_URL,
            "embeddings_ready": embeddings_ready
        })
    except Exception as e:
        print(f"Error checking NIM status: {str(e)}")
        return jsonify({
            "connected": False,
            "status": "error",
            "error": str(e)
        }), 500

@app.route('/api/debug', methods=['GET'])
def debug_info():
    """Diagnostic endpoint to check system status"""
    global embeddings_ready, docstore, convstore

    # Check if embeddings are actually initialized
    embedding_status = {
        "embeddings_ready_flag": embeddings_ready,
        "docstore_exists": docstore is not None,
        "convstore_exists": convstore is not None,
    }

    # Get counts of documents if possible
    if docstore is not None:
        try:
            embedding_status["docstore_index_ntotal"] = docstore.index.ntotal
        except:
            embedding_status["docstore_index_ntotal"] = "error accessing"

    # Testing retrieval on a simple query if docstore exists
    retrieval_test = {}
    if docstore is not None:
        try:
            results = docstore.similarity_search_with_score("Jetson", k=1)
            retrieval_test["success"] = True
            retrieval_test["result_count"] = len(results)
            if results:
                doc, score = results[0]
                retrieval_test["first_result"] = {
                    "content_preview": doc.page_content[:100] + "...",
                    "score": float(score),
                    "metadata": doc.metadata
                }
        except Exception as e:
            retrieval_test["success"] = False
            retrieval_test["error"] = str(e)

    # Testing the LLM directly
    llm_test = {}
    try:
        response = instruct_llm.invoke("Say hello")
        llm_test["success"] = True
        llm_test["response"] = response.content
    except Exception as e:
        llm_test["success"] = False
        llm_test["error"] = str(e)

    return jsonify({
        "server_status": "running",
        "embeddings": embedding_status,
        "retrieval_test": retrieval_test,
        "llm_test": llm_test,
        "api_key_configured": bool(NVIDIA_API_KEY),
        "model": MODEL
    })

@app.route('/UI/NVLogo_2D.jpg')
def serve_nvidia_logo():
    """Special route to serve the NVIDIA logo"""
    try:
        # Try sending from current directory
        ui_path = os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))), 'UI')
        logo_path = os.path.join(ui_path, 'NVLogo_2D.jpg')

        if os.path.exists(logo_path):
            print(f"NVIDIA logo found at: {logo_path}")
            return send_from_directory('../UI', 'NVLogo_2D.jpg')
        else:
            print(f"NVIDIA logo not found at: {logo_path}")
            # Return a default image or error message
            return "Logo not found", 404
    except Exception as e:
        print(f"Error serving NVIDIA logo: {str(e)}")
        return "Error serving logo", 500

if __name__ == '__main__':
    print("Starting Flask server...")
    # Use a more reliable server configuration
    app.run(debug=False, threaded=True, port=3000)